#!/usr/bin/env python3
"""
Генератор щоденника розробки дистрибутива Devuan Security OS
Форматування відповідно до ДСТУ 3008:2015
Запуск: PYTHONPATH=~/.nix-profile/lib/python3.13/site-packages python3 generate_diary.py
"""

import os
import sys
import glob

nix_site_packages = os.path.expanduser("~/.nix-profile/lib/python3.13/site-packages")
if nix_site_packages not in sys.path:
    sys.path.insert(0, nix_site_packages)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# === КОНФІГУРАЦІЯ ===
SCREENSHOTS_DIR = "/home/zoozienix/projects/Devuan_Project/screenshots"
OUTPUT_FILE = "/home/zoozienix/projects/Devuan_Project/diary/Devuan_Security_OS_Diary.docx"

# ДСТУ 3008:2015 — параметри сторінки
PAGE_LEFT   = Cm(2.5)
PAGE_RIGHT  = Cm(1.5)
PAGE_TOP    = Cm(2.0)
PAGE_BOTTOM = Cm(2.0)

# Шрифт
FONT_MAIN  = "Times New Roman"
FONT_SIZE  = Pt(14)
FONT_CODE  = "Courier New"
INDENT     = Cm(1.25)  # відступ першого рядка абзацу

# ============================================================
# ДОПОМІЖНІ ФУНКЦІЇ
# ============================================================

def apply_dstu_paragraph(paragraph, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Застосовує базове ДСТУ-форматування до абзацу"""
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = INDENT
    else:
        pf.first_line_indent = Pt(0)

def apply_dstu_run(run, bold=False, italic=False, size=None):
    """Застосовує ДСТУ-форматування до run"""
    run.font.name = FONT_MAIN
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_dstu_paragraph(doc, text="", bold=False, italic=False,
                        first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Додає абзац у стилі ДСТУ"""
    p = doc.add_paragraph()
    apply_dstu_paragraph(p, first_line_indent, align)
    if text:
        run = p.add_run(text)
        apply_dstu_run(run, bold=bold, italic=italic)
    return p

def add_dstu_heading(doc, text, level=1):
    """
    Заголовки за ДСТУ 3008:2015:
      level=1 → розділ (14pt, жирний, по центру, велика літера)
      level=2 → підрозділ (14pt, жирний, з відступом)
      level=3 → пункт (14pt, жирний курсив, з відступом)
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)

    run = p.add_run(text)
    run.font.name = FONT_MAIN
    run.font.size = FONT_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)

    if level == 1:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        run.font.all_caps = True
    elif level == 2:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = INDENT
        run.bold = True
    elif level == 3:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = INDENT
        run.bold = True
        run.italic = True
    return p

def add_screenshot(doc, filepath, fig_num, caption_text, width_cm=14):
    """Вставляє рисунок з підписом у стилі ДСТУ"""
    add_dstu_paragraph(doc, first_line_indent=False)  # відступ перед рисунком

    # Пошук файлу (на випадок пробілів у назві)
    actual_path = filepath
    if not os.path.exists(filepath):
        matches = glob.glob(filepath.rstrip() + "*")
        if matches:
            actual_path = matches[0]

    if os.path.exists(actual_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        run = p_img.add_run()
        run.add_picture(actual_path, width=Cm(width_cm))
    else:
        p_err = add_dstu_paragraph(
            doc, f"[РИСУНОК ВІДСУТНІЙ: {os.path.basename(filepath)}]",
            align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False
        )
        p_err.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Підпис рисунку: "Рисунок N — Назва" (по центру, без відступу)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after = Pt(6)
    p_cap.paragraph_format.first_line_indent = Pt(0)
    p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p_cap.add_run(f"Рисунок {fig_num} \u2014 {caption_text}")
    run.font.name = FONT_MAIN
    run.font.size = FONT_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_dstu_table(doc, headers, rows):
    """Додає таблицю у стилі ДСТУ"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Заголовки
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], "C0C0C0")
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(h)
        run.font.name = FONT_MAIN
        run.font.size = FONT_SIZE
        run.font.bold = True

    # Рядки
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(val)
            run.font.name = FONT_MAIN
            run.font.size = FONT_SIZE
            if i == 0:
                run.font.bold = True

    return table

def add_page_number(doc):
    """Додає нумерацію сторінок знизу по центру"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.font.name = FONT_MAIN
    run.font.size = FONT_SIZE

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ============================================================
# ГЕНЕРАЦІЯ ДОКУМЕНТУ
# ============================================================

def generate():
    doc = Document()

    # --- Параметри сторінки (ДСТУ 3008:2015) ---
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = PAGE_LEFT
    section.right_margin  = PAGE_RIGHT
    section.top_margin    = PAGE_TOP
    section.bottom_margin = PAGE_BOTTOM

    # Нумерація сторінок
    add_page_number(doc)

    # Очищаємо стиль Normal
    normal = doc.styles['Normal']
    normal.font.name = FONT_MAIN
    normal.font.size = FONT_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ==========================================
    # ТИТУЛЬНА СТОРІНКА
    # ==========================================
    for _ in range(3):
        add_dstu_paragraph(doc, first_line_indent=False)

    # Назва організації
    add_dstu_paragraph(
        doc, "Щоденник розробки дипломного проєкту",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, bold=True
    )
    add_dstu_paragraph(doc, first_line_indent=False)

    # Назва проєкту
    p_title = add_dstu_paragraph(
        doc, "DEVUAN SECURITY OS",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, bold=True
    )
    p_title.runs[0].font.size = Pt(16)

    add_dstu_paragraph(doc, first_line_indent=False)
    add_dstu_paragraph(
        doc,
        "Кастомний дистрибутив GNU/Linux на базі Devuan Excalibur 6.x,\n"
        "орієнтований на безпеку, без системи ініціалізації systemd",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, italic=True
    )

    for _ in range(8):
        add_dstu_paragraph(doc, first_line_indent=False)

    meta = [
        ("Автор:", "zoozienix"),
        ("Репозиторій:", "github.com/zoozieuniver/Devuan_Project"),
        ("Дата початку:", datetime.date.today().strftime("%d.%m.%Y")),
        ("Версія документу:", "0.1"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.first_line_indent = Pt(0)
        r1 = p.add_run(f"{label} ")
        apply_dstu_run(r1, bold=True)
        r2 = p.add_run(value)
        apply_dstu_run(r2)

    doc.add_page_break()

    # ==========================================
    # ЗМІСТ
    # ==========================================
    add_dstu_heading(doc, "ЗМІСТ", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    toc_items = [
        ("Вступ", "3"),
        ("1 Технічний стек та обґрунтування вибору", "3"),
        ("2 Фаза 1. Підготовка середовища (Build Box)", "4"),
        ("  2.1 Завантаження ISO-образу Devuan", "4"),
        ("  2.2 Створення віртуальної машини у VMware", "5"),
        ("3 Фаза 2. Встановлення Devuan без графічного середовища", "\u2014"),
        ("4 Фаза 3. Конфігурування live-build", "\u2014"),
        ("5 Фаза 4. Збірка ISO-образу", "\u2014"),
        ("6 Фаза 5. Тестування та оптимізація", "\u2014"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item)
        apply_dstu_run(r, bold=item[0].isdigit() and "  " not in item)
        dots = "." * max(2, 65 - len(item))
        r2 = p.add_run(f" {dots} {page}")
        apply_dstu_run(r2)

    doc.add_page_break()

    # ==========================================
    # ВСТУП
    # ==========================================
    add_dstu_heading(doc, "ВСТУП", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Даний щоденник фіксує процес розробки власного дистрибутиву GNU/Linux на базі "
        "Devuan Excalibur 6.x. Devuan є форком дистрибутиву Debian, головною особливістю "
        "якого є відмова від системи ініціалізації systemd на користь класичного SysVinit."
    )
    add_dstu_paragraph(
        doc,
        "Метою проєкту є створення мінімальної, безпечної та відтворюваної (reproducible) "
        "системи, яка може використовуватися як основа для спеціалізованих середовищ безпеки. "
        "Кінцевим результатом є завантажувальний ISO-образ, зібраний інструментом live-build, "
        "з повністю задокументованим та версіонованим процесом збірки через систему контролю "
        "версій Git."
    )
    add_dstu_paragraph(
        doc,
        "Кожен крок розробки фіксується у вигляді Git-комітів із відповідними скріншотами, "
        "що забезпечує повну відтворюваність проєкту та слугує технічним щоденником."
    )

    doc.add_page_break()

    # ==========================================
    # РОЗДІЛ 1: ТЕХНІЧНИЙ СТЕК
    # ==========================================
    add_dstu_heading(doc, "1 ТЕХНІЧНИЙ СТЕК ТА ОБҐРУНТУВАННЯ ВИБОРУ", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "У таблиці 1.1 наведено перелік технологій та інструментів, обраних для реалізації "
        "проєкту, із зазначенням обґрунтування кожного вибору."
    )
    add_dstu_paragraph(doc, "Таблиця 1.1 \u2014 Технічний стек проєкту",
                       align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

    add_dstu_table(doc,
        ["Компонент", "Вибір", "Обґрунтування"],
        [
            ("База дистрибутиву", "Devuan Excalibur 6.1.1", "Debian без systemd, стабільна база"),
            ("Ядро", "Linux (стабільний реліз)", "Широка підтримка обладнання"),
            ("Init система", "SysVinit", "Простота, передбачуваність, без залежностей"),
            ("Робочий стіл", "XFCE 4", "Легковаговість, повна функціональність"),
            ("Інструмент збірки", "live-build", "Декларативна конфігурація, стандарт Debian"),
            ("Гіпервізор", "VMware Workstation Pro 25H2", "Build Box та тестове середовище"),
            ("Версіонування", "Git + GitHub", "Відтворюваність збірки, журнал змін"),
        ]
    )

    doc.add_page_break()

    # ==========================================
    # РОЗДІЛ 2: ФАЗА 1
    # ==========================================
    add_dstu_heading(doc, "2 ФАЗА 1. ПІДГОТОВКА СЕРЕДОВИЩА (BUILD BOX)", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Метою даної фази є підготовка еталонного середовища збірки (Build Box) у вигляді "
        "віртуальної машини під управлінням VMware Workstation Pro 25H2. Саме в цьому "
        "середовищі надалі виконуватиметься збірка дистрибутиву за допомогою live-build."
    )

    # 2.1
    add_dstu_heading(doc, "2.1 Завантаження ISO-образу Devuan", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Для встановлення Build Box обрано образ типу netinstall дистрибутиву "
        "Devuan Excalibur 6.1.1 (amd64). Даний тип образу є мінімальним встановлювачем, "
        "що завантажує лише базову систему, а всі пакети отримує з мережі під час "
        "встановлення. Це забезпечує актуальність встановлених пакетів."
    )
    add_dstu_paragraph(
        doc,
        "Образ завантажено з офіційного архіву Devuan за адресою files.devuan.org. "
        "Версія Excalibur 6.x базується на Debian 13 (Trixie) та є поточним стабільним "
        "релізом станом на липень 2026 року. Розмір образу складає 594 МБ."
    )
    add_dstu_paragraph(doc, first_line_indent=False)

    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "1.1_devuan_release_archive.png"),
        "2.1", "Офіційний архів ISO-образів Devuan Excalibur на files.devuan.org"
    )

    # 2.2
    add_dstu_heading(doc, "2.2 Створення віртуальної машини у VMware Workstation", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Для Build Box створено нову віртуальну машину у VMware Workstation Pro 25H2. "
        "Параметри підібрані з урахуванням вимог процесу збірки через live-build, "
        "що потребує достатньо оперативної пам'яті та обчислювальних ресурсів для "
        "завантаження та розпакування пакетів."
    )

    add_dstu_paragraph(doc, "У таблиці 2.1 наведено параметри створеної віртуальної машини.")
    add_dstu_paragraph(doc, "Таблиця 2.1 \u2014 Параметри віртуальної машини Build Box",
                       align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

    add_dstu_table(doc,
        ["Параметр", "Значення"],
        [
            ("Назва ВМ", "Devuan_Security"),
            ("Гостьова ОС (VMware)", "Debian 13.x 64-bit"),
            ("Тип конфігурації", "Typical (Recommended)"),
            ("ISO-образ", "devuan_excalibur_6.1.1_amd64_netinstall.iso"),
            ("Оперативна пам'ять", "2048 МБ (2 ГБ)"),
            ("Процесор", "2 ядра"),
            ("Жорсткий диск", "20 ГБ (Split into multiple files)"),
            ("Мережевий адаптер", "NAT"),
        ]
    )

    add_dstu_paragraph(doc, first_line_indent=False)
    add_dstu_paragraph(
        doc,
        "На рисунках 2.2\u20132.8 наведено послідовність екранів майстра створення "
        "нової віртуальної машини у VMware Workstation Pro 25H2."
    )

    screenshots_2 = [
        ("1.2_vmware_main_screen.png",      "2.2", "Головне меню VMware Workstation Pro 25H2"),
        ("1.3_vmware_wizard_type.png",      "2.3", "Вибір типу конфігурації ВМ (Typical / Custom)"),
        ("1.4_vmware_iso_selection.png",    "2.4", "Підключення ISO-образу Devuan Excalibur 6.1.1"),
        ("1.5_vmware_guest_os_debian13.png","2.5", "Вибір гостьової ОС: Debian 13.x 64-bit"),
        ("1.6_vmware_vm_name.png",          "2.6", "Задання імені віртуальної машини: Devuan_Security"),
        ("1.7_vmware_disk_size.png",        "2.7", "Налаштування розміру віртуального диску: 20 ГБ"),
        ("1.8_vmware_summary.png",          "2.8", "Фінальні параметри ВМ перед створенням"),
    ]
    for fname, num, caption in screenshots_2:
        add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, fname), num, caption)

    # ==========================================
    # РОЗДІЛ 3–6: ЗАГЛУШКИ
    # ==========================================
    # ==========================================
    # РОЗДІЛ 3: ФАЗА 2
    # ==========================================
    doc.add_page_break()
    add_dstu_heading(doc, "3 ФАЗА 2. ВСТАНОВЛЕННЯ DEVUAN БЕЗ ГРАФІЧНОГО СЕРЕДОВИЩА", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Метою цієї фази є безпосереднє встановлення операційної системи Devuan Excalibur 6.1.1 "
        "на віртуальний диск створеної ВМ. Для мінімізації споживання ресурсів та підвищення "
        "безпеки системи, встановлення проводиться у мінімальній конфігурації (без графічної оболонки)."
    )

    # 3.1
    add_dstu_heading(doc, "3.1 Початкове завантаження та вибір локалізації", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "При завантаженні з інсталяційного образу користувачу пропонується menu вибору режиму "
        "встановлення. Обрано стандартний текстовий режим встановлення (Install). На наступних "
        "кроках налаштовано мову інтерфейсу, країну місцезнаходження та розкладку клавіатури."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.1_devuan_boot_menu.png"), "3.1", "Меню завантаження Devuan Excalibur")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.2_devuan_language.png"), "3.2", "Вибір мови встановлення (English)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.3.1_devuan_location_other.png"), "3.3", "Вибір регіону місцезнаходження (Other)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.3.2_devuan_location_europe.png"), "3.4", "Вибір континенту (Europe)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.3.3_devuan_location_ukraine.png"), "3.5", "Вибір країни (Ukraine)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.4_devuan_locale.png"), "3.6", "Вибір локалі за замовчуванням (en_US.UTF-8)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.5_devuan_keyboard.png"), "3.7", "Вибір розкладки клавіатури (American English)")

    # 3.2
    add_dstu_heading(doc, "3.2 Налаштування мережі та облікових записів", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "На цьому етапі відбувається автоматичне отримання мережевих налаштувань через DHCP, "
        "введення імені хоста (hostname), а також створення облікових записів адміністратора (root) "
        "та звичайного користувача системи."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.6_devuan_hostname.png"), "3.8", "Задання імені хоста (devuan-security)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.7_devuan_root_password.png"), "3.9", "Встановлення паролю суперкористувача (root)")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.8_devuan_user_create.png"), "3.10", "Створення облікового запису звичайного користувача")

    # 3.3
    add_dstu_heading(doc, "3.3 Розмітка дискового простору", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Для встановлення системи обрано метод автоматичної розмітки всього дискового простору (Guided - use entire disk). "
        "Оскільки створюється віртуальна машина розробника (Build Box), обрано найпростішу схему розмітки, де всі файли "
        "розташовуються на одному розділі (All files in one partition)."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.10.1_devuan_disk_method.png"), "3.11", "Вибір методу розмітки диска")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.10.2_devuan_disk_select.png"), "3.12", "Вибір цільового диска для розмітки")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.10.3_devuan_disk_scheme.png"), "3.13", "Вибір схеми розмітки диска")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.10.4_devuan_disk_overview.png"), "3.14", "Огляд створених розділів")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.10.5_devuan_disk_confirm.png"), "3.15", "Підтвердження запису таблиці розділів на диск")

    # 3.4
    add_dstu_heading(doc, "3.4 Налаштування пакетного менеджера та вибір ПЗ", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Після копіювання базової системи виконується конфігурування джерела пакунків (APT mirror) та вибір "
        "програмного забезпечення для встановлення. У вікні вибору ПЗ (Tasksel) відключено всі графічні "
        "середовища, залишено лише стандартні системні утиліти. Також обрано систему ініціалізації sysvinit."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.11.1_devuan_mirror_country.png"), "3.16", "Вибір країни дзеркала репозиторіїв")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.11.2_devuan_mirror_select.png"), "3.17", "Вибір адреси дзеркала репозиторію")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.11.3_devuan_software_selection.png"), "3.18", "Вибір компонентів системи у меню Tasksel")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.11.4_devuan_init_system.png"), "3.19", "Вибір системи ініціалізації (sysvinit)")

    # 3.5
    add_dstu_heading(doc, "3.5 Встановлення завантажувача та перший вхід", level=2)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_paragraph(
        doc,
        "Завершальним етапом є встановлення системного завантажувача GRUB у головний завантажувальний запис (MBR) "
        "диска /dev/sda та перезавантаження системи для першого запуску."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.12.1_devuan_grub_install.png"), "3.20", "Запит на встановлення завантажувача GRUB")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.12.2_devuan_grub_device.png"), "3.21", "Вибір пристрою для встановлення завантажувача GRUB")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.12_devuan_install_complete.png"), "3.22", "Повідомлення про успішне завершення встановлення")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.13_devuan_grub_menu.png"), "3.23", "Стартове меню завантажувача GRUB встановленої системи")
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "2.14_devuan_first_login.png"), "3.24", "Перший вхід до консолі встановленої ОС")

    # ==========================================
    # ФАЗА 3. КОНФІГУРУВАННЯ LIVE-BUILD
    # ==========================================
    doc.add_page_break()
    add_dstu_heading(doc, "4 ФАЗА 3. КОНФІГУРУВАННЯ LIVE-BUILD", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_heading(doc, "4.1 Принципи та структура конфігурації", level=2)
    add_dstu_paragraph(
        doc,
        "Для побудови дистрибутиву Devuan Security OS використовується інструментарій live-build, який дозволяє повністю "
        "автоматизувати процес створення live-систем на базі Debian-подібних ОС. Конфігурація системи описується за допомогою "
        "декларативних списків пакетів, скриптів автоматизації та конфігураційних файлів, які переносяться у складальне "
        "середовище (Build Box) та використовуються під час збірки."
    )

    add_dstu_heading(doc, "4.2 Сценарій ініціалізації auto/config", level=2)
    add_dstu_paragraph(
        doc,
        "Сценарій auto/config використовується для генерації базового дерева конфігурації. У ньому визначаються ключові "
        "параметри дистрибутиву: робочий режим встановлюється в значення devuan, цільовий дистрибутив — excalibur (Devuan 6), "
        "архітектура процесора — amd64, тип вихідного файлу — iso-hybrid. Для оптимізації процесу збірки та уникнення конфліктів "
        "із застарілими серверами оновлень Debian, параметр security вимкнено (оскільки в Devuan оновлення безпеки інтегровано в "
        "основний репозиторій). Програма налаштовується на використання дзеркала репозиторіїв http://deb.devuan.org/merged."
    )

    add_dstu_heading(doc, "4.3 Списки пакетів та кастомізація середовища", level=2)
    add_dstu_paragraph(
        doc,
        "Списки пакетів для встановлення описуються у файлі config/package-lists/custom.list.chroot. Конфігурація містить: "
        "графічну оболонку XFCE 4 (пакети xfce4, xfce4-goodies) та дисплейний менеджер lightdm; базовий сервер відображення xorg; "
        "пакети безпеки та аудиту мережі (ufw, nmap, wireshark, tshark, tcpdump, keepassxc, macchanger, tor); службові утиліти "
        "(sudo, htop, curl, wget, git, fastfetch, bash-completion). Через видалення застарілої утиліти neofetch із репозиторіїв "
        "Devuan Excalibur, у збірці застосовано сучасну альтернативу fastfetch."
    )
    add_dstu_paragraph(
        doc,
        "Для конфігурування мережі у live-системі створено файл config/includes.chroot/etc/network/interfaces, який забезпечує "
        "автоматичне підняття інтерфейсів eth0 та wlan0 за допомогою DHCP. Також додано конфігурацію config/includes.chroot/etc/sudoers.d/live, "
        "яка надає стандартному live-користувачу (user) права виконання адміністративних команд без запиту пароля. Додатково "
        "виконано патч скрипту /usr/lib/live/build/config у ВМ для примусового вибору підсистеми live-config-sysvinit (замість systemd) "
        "при використанні ініціалізації sysvinit."
    )

    # ==========================================
    # ФАЗА 4. ЗБІРКА ISO-ОБРАЗУ
    # ==========================================
    doc.add_page_break()
    add_dstu_heading(doc, "5 ФАЗА 4. ЗБІРКА ISO-ОБРАЗУ", level=1)
    add_dstu_paragraph(doc, first_line_indent=False)

    add_dstu_heading(doc, "5.1 Процес компонування та створення SquashFS", level=2)
    add_dstu_paragraph(
        doc,
        "Процес збірки запускається послідовністю команд lb clean, lb config та sudo lb build. Під час виконання відбувається "
        "створення базової кореневої файлової системи (bootstrap), встановлення ядра Linux, системного завантажувача та всіх "
        "заданих у конфігурації пакетів. На завершальному етапі вміст chroot-оточення пакується у стиснуту файлову систему SquashFS."
    )
    add_dstu_paragraph(
        doc,
        "На рисунку 3.1 зображено процес пакування образу файлової системи за допомогою squashfs-tools у складальному "
        "середовищі віртуальної машини."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "3.1_devuan_building_process.png"), "3.25", "Створення стиснутого образу файлової системи SquashFS")

    add_dstu_heading(doc, "5.2 Генерація гібридного ISO-образу та його трансфер", level=2)
    add_dstu_paragraph(
        doc,
        "Після завершення створення SquashFS інструментарій live-build інтегрує його із завантажувачем syslinux та створює "
        "фінальний гібридний образ devuan-security-amd64-amd64.hybrid.iso. Отриманий файл розміром 1018 МБ копіюється з "
        "віртуальної машини на хост-систему за допомогою безпечного протоколу SCP."
    )
    add_dstu_paragraph(
        doc,
        "На рисунку 3.2 показано результат виведення розміру та прав доступу створеного ISO-образу на хост-машині."
    )

    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, "3.2_devuan_iso_created.png"), "3.26", "Результат створення гібридного ISO-образу на хост-машині")

    # Placeholder for Phase 5
    placeholders = [
        ("6 ФАЗА 5. ТЕСТУВАННЯ ТА ОПТИМІЗАЦІЯ (QA)", "", []),
    ]

    for title, body, items in placeholders:
        doc.add_page_break()
        add_dstu_heading(doc, title, level=1)
        add_dstu_paragraph(doc, first_line_indent=False)
        add_dstu_paragraph(doc,
            "Розділ буде заповнено на відповідному етапі розробки (тестування live-образу у ВМ).", italic=True)

    # ==========================================
    # ЗБЕРЕЖЕННЯ
    # ==========================================
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"✅ Щоденник (ДСТУ) збережено: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate()
